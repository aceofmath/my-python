from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openpyxl import load_workbook
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc_template = Document(r'C:\dev\workspace\py_test\resource\docx\수료증.docx')
wb = load_workbook(r'C:\dev\workspace\py_test\resource\xlsx\수강생명단.xlsx')
ws = wb.active

# 오늘 날짜를 가져옵니다.
today = datetime.date.today()

for i, paragraph in  enumerate(doc_template.paragraphs):
    print(str(i)  +  ": "  + paragraph.text)

for i in range(1, ws.max_row + 1):
    if i == 1:  # 헤더 스킵
        continue

    name = ws.cell(row=i, column=1).value
    birth = ws.cell(row=i, column=2).value
    period = ws.cell(row=i, column=4).value

    # 원본 문서 복사해서 사용
    doc = Document(r'C:\dev\workspace\py_test\resource\docx\수료증.docx')

    # 성명 수정
    p = doc.paragraphs[1]
    p.clear()  # 기존 run 삭제
    run = p.add_run(f"성 명: {name}")
    run.font.name = '맑은 고딕'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 생년월일 수정
    p = doc.paragraphs[2]
    p.clear()
    run = p.add_run(f"생년월일: {str(birth)[:10]}")
    run.font.name = '맑은 고딕'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 수료기간 수정
    p = doc.paragraphs[3]
    p.clear()
    run = p.add_run(f"수료기간: {period}")
    run.font.name = '맑은 고딕'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 발행일자 수정
    p = doc.paragraphs[13]
    p.clear()
    run = p.add_run(f"{today.strftime("%Y년 %m월 %d일")}")
    run.font.name = '맑은 고딕'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    table = doc.tables[0]  # 첫 번째 테이블 기준 (필요 시 인덱스 변경)
    # 예: 1행 0열에 성명, 2행 0열에 생년월일이 있다고 가정
    name_cell = table.cell(0, 1)
    name_cell.text = ""  # 기존 내용 삭제
    run = name_cell.paragraphs[0].add_run(f"(주)에이치쓰리연구소")
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(12)
    name_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 가운데 정렬

    # 저장
    doc.save(rf'C:\dev\workspace\py_test\result\수료증_{name}.docx')
