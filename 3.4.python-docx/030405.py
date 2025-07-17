from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

doc = Document(r'C:\dev\workspace\py_test\resource\docx\예제 문서.docx')

# ======================================================
# 1. 문단 정렬
# ======================================================

# 왼쪽 정렬
paragraph1 = doc.paragraphs[1]
paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT

# 가운데 정렬
paragraph2 = doc.paragraphs[2]
paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 오른쪽 정렬
paragraph3 = doc.paragraphs[3]
paragraph3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 양쪽 정렬
paragraph4 = doc.paragraphs[4]
paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# 텍스트 배분 (글자를 흩어서 배치)
paragraph_last = doc.paragraphs[-1]  # 마지막 문단
paragraph_last.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE

# ======================================================
# 2. 셀 정렬
# ======================================================

# 1) 수평 정렬
# LEFT : 왼쪽 정렬, CENTER: 가운데 정렬, RIGHT: 오른쪽 정렬
doc.tables[0].rows[0].cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
doc.tables[0].rows[0].cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
doc.tables[0].rows[0].cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT

# 2) 수직 정렬
# LEFT : 위쪽 정렬, CENTER: 가운데 정렬, RIGHT: 아래쪽 정렬
doc.tables[0].rows[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
doc.tables[0].rows[0].cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
doc.tables[0].rows[0].cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

# 현재 작업경로에 저장
doc.save(r'C:\dev\workspace\py_test\result\030405.docx')
