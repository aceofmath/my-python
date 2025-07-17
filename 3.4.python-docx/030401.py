# 가장 기본적인 기능(문서 열기, 저장, 글자 쓰기 등등)
from docx import Document

# 문단 정렬
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 문자 스타일 변경
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# 제목
doc.add_heading('가장 큰 제목 (아래에 밑줄)', level=0)
doc.add_heading('제목 크기, H1', level=1)
doc.add_heading('제목 크기, H2', level=2)
doc.add_heading('제목 크기, H3', level=3)
doc.add_heading('제목 크기, H4', level=4)
doc.add_heading('제목 크기, H5', level=5)
doc.add_heading('제목 크기, H6', level=6)

p = doc.add_paragraph('두번째 문단: 여기에 원하는 텍스트를 마음껏 입력하면 됩니다.')

# 굵은 글씨(Bold) 적용
p.add_run('문단에 굵은 글자 추가').bold = True

# 기울임꼴(Italic) 적용
p.add_run('문단에 이텔릭 글자 추가').italic = True

# 밑줄(Underline) 적용
p.add_run('문단에 밑줄 글자 추가').underline = True


# 현재 작업경로에 저장
doc.save(r'C:\dev\workspace\py_test\result\030401.docx')
