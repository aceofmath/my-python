from docx import Document

# Cm와 Inch 단위를 사용하기 위한 모듈
from docx.shared import Cm, Inches

doc = Document()

# ======================================================
# 1. 이미지 삽입
# ======================================================

# 사진의 크기를 Cm 단위로 설정하여 삽입
doc.add_picture(r'C:\dev\workspace\py_test\resource\img\cat.jpg',width= Cm(16), height= Cm(9))

# 사진의 크기를 Inch 단위로 설정하여 삽입
# doc.add_picture(r'C:\dev\workspace\py_test\resource\img\cat.jpg',width= Inches(4), height= Inches(3))


# ======================================================
# 2. 표 삽입
# ======================================================

# 1) 행과 열 설정 및 표 만들기
# 2행 3열의 표 만들기
table = doc.add_table(rows = 2, cols = 3)

# 만든 표의 스타일을 가장 기본 스타일인 'Table Grid'로 설정
table.style = doc.styles['Table Grid']


# 표의 첫 행을 리스트로 가져오기
first_row = table.rows[0].cells

# 2) 각 셀에 내용 입력
# 첫 행의 각 열들에 접근해서 값 입력 
first_row[0].text = 'a'
first_row[1].text = 'b'
first_row[2].text = 'c'

# 표의 두번째 행을 리스트로 가져온 후, 각 셀에 값 입력
second_row = table.rows[1].cells
second_row[0].text = 'd'
second_row[1].text = 'e'
second_row[2].text = 'f'

# 3) 행과 열 추가하기
row_cells = table.add_row().cells
col_cells = table.add_column(width=Cm(2)).cells

# 저장
doc.save(r'C:\dev\workspace\py_test\result\030402.docx')
