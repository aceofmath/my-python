from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor


doc = Document()

# ======================================================
# 2. 크기
# ======================================================
para = doc.add_paragraph('이 글자의 크기를 바꿔봅시다')

# 첫번째 문단의 문장(run)들을 리스트로 받기
para = doc.paragraphs[0].runs

# for 문을 이용해서 
for run in para:
    run.font.size = Pt(20)

# ======================================================
# 3. 색깔
# ======================================================

para2 = doc.add_paragraph('글자 색깔을 바꿔봅시다')
run = para2.runs[0]
font = run.font

# RGB 컬러를 각각 16진수로 표현 (R, G, B)
font.color.rgb = RGBColor(0xFF, 0x24, 0xE9)

# 저장
doc.save(r'C:\dev\workspace\py_test\result\030404.docx')
