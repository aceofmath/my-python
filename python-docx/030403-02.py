from docx import Document

doc = Document(r'C:\dev\workspace\py_test\예제 문서.docx')

# ======================================================
# 1. paragraph 인덱싱
# ======================================================

p = doc.paragraphs[4]
p.add_run('\n문단에 글자 추가')

doc.paragraphs[3].insert_paragraph_before("문장을 삽입한다.")



# ======================================================
# 1. paragraph 인덱싱
# ======================================================

for i, paragraph in enumerate(doc.paragraphs):
    print(str(i+1) + ": " + paragraph.text)


# ======================================================
# 2. table 인덱싱
# ======================================================

# 문서 안의 모든 표를 가져옴
# tables = doc.tables

# 가장 처음 표의 첫행, 첫열의 첫문단 내용 가져오기
# tables[0].rows[0].cells[0].paragraphs[0].text

table = doc.tables[0]

for row in table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            print(para.text)


# # 첫번째 표 인덱싱
# table = doc.tables[0]

# for row in table.rows:
#     for cell in row.cells:
#         for para in cell.paragraphs:
#             if(para.text == "하나"):
#                 para.add_run(' <-- 찾았다 하나')

# # 확인을 위해서 저장하기
# doc.save('예제 문서.docx')